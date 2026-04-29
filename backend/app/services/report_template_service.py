import base64
import json
import mimetypes
import re
import zipfile
from datetime import datetime
from io import BytesIO
from pathlib import Path
from xml.etree import ElementTree

from fastapi import UploadFile
from sqlalchemy.orm import Session

from app.core.config import get_settings
from app.models.report_template import ReportTemplate
from app.repositories.report_template_repository import ReportTemplateRepository
from app.schemas.report_template import (
    ReportTemplateCreate,
    ReportTemplateReferenceAnalysis,
    ReportTemplateReferenceText,
    ReportTemplateUpdate,
)
from app.services.report_service import _generate_claude, _generate_gemini, _generate_openai, _report_provider_order
from app.services.settings_service import get_effective_provider_settings


MAX_REFERENCE_FILE_BYTES = 15 * 1024 * 1024
TEXT_REFERENCE_EXTENSIONS = {".txt", ".md", ".markdown", ".csv"}
DOCX_REFERENCE_EXTENSIONS = {".docx"}
ODT_REFERENCE_EXTENSIONS = {".odt"}
PDF_REFERENCE_EXTENSIONS = {".pdf"}
IMAGE_REFERENCE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp"}


def get_template(db: Session, template_id: str, workspace_id: str = "local-workspace") -> ReportTemplate:
    template = ReportTemplateRepository(db).get_for_workspace(template_id, workspace_id)
    if not template:
        raise ValueError("Modelo não encontrado")
    return template


def _build_duplicate_name(repository: ReportTemplateRepository, original_name: str, workspace_id: str) -> str:
    base_name = f"{original_name} (cópia)"
    if not repository.get_by_name(base_name, workspace_id):
        return base_name

    index = 2
    while repository.get_by_name(f"{base_name} {index}", workspace_id):
        index += 1
    return f"{base_name} {index}"


def list_templates(db: Session, workspace_id: str = "local-workspace") -> list[ReportTemplate]:
    return ReportTemplateRepository(db).list(workspace_id)


def create_template(db: Session, payload: ReportTemplateCreate, workspace_id: str = "local-workspace") -> ReportTemplate:
    repository = ReportTemplateRepository(db)
    if repository.get_by_name(payload.name, workspace_id):
        raise ValueError("Já existe um modelo com esse nome")
    return repository.create(ReportTemplate(workspace_id=workspace_id, **payload.model_dump()))


def _read_reference_file(reference_file: UploadFile) -> bytes:
    data = reference_file.file.read()
    if not data:
        raise ValueError("Arquivo de referencia vazio")
    if len(data) > MAX_REFERENCE_FILE_BYTES:
        raise ValueError("Arquivo de referencia excede o limite de 15 MB")
    return data


def _decode_text(data: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "latin-1"):
        try:
            return data.decode(encoding).strip()
        except UnicodeDecodeError:
            continue
    raise ValueError("Nao foi possivel ler o texto do arquivo")


def _extract_docx_text(data: bytes) -> str:
    from docx import Document

    document = Document(BytesIO(data))
    blocks: list[str] = []
    blocks.extend(paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip())
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                blocks.append(" | ".join(cells))
    return "\n".join(blocks).strip()


def _extract_odt_text(data: bytes) -> str:
    try:
        with zipfile.ZipFile(BytesIO(data)) as archive:
            content = archive.read("content.xml")
    except (KeyError, zipfile.BadZipFile) as exc:
        raise ValueError("Nao foi possivel ler o arquivo ODT.") from exc

    root = ElementTree.fromstring(content)
    paragraphs: list[str] = []
    for element in root.iter():
        if not element.tag.endswith(("}p", "}h")):
            continue
        text = "".join(element.itertext()).strip()
        if text:
            paragraphs.append(text)
    return "\n".join(paragraphs).strip()


def _extract_pdf_text(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise ValueError("Para analisar PDF, instale a dependencia pypdf.") from exc

    reader = PdfReader(BytesIO(data))
    pages = [(page.extract_text() or "").strip() for page in reader.pages]
    return "\n\n".join(page for page in pages if page).strip()


def _is_image_reference(filename: str, content_type: str | None) -> bool:
    suffix = Path(filename).suffix.lower()
    return suffix in IMAGE_REFERENCE_EXTENSIONS or bool(content_type and content_type.startswith("image/"))


def _extract_reference_text(filename: str, content_type: str | None, data: bytes) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix in DOCX_REFERENCE_EXTENSIONS:
        text = _extract_docx_text(data)
    elif suffix in ODT_REFERENCE_EXTENSIONS:
        text = _extract_odt_text(data)
    elif suffix in PDF_REFERENCE_EXTENSIONS:
        text = _extract_pdf_text(data)
    elif suffix in TEXT_REFERENCE_EXTENSIONS or (content_type and content_type.startswith("text/")):
        text = _decode_text(data)
    else:
        text = _decode_text(data)

    if not text:
        raise ValueError("Nao foi possivel extrair texto do documento")
    return text


def _build_docx_base64_from_text(title: str, content: str) -> str:
    from docx import Document

    document = Document()
    document.add_heading(title, level=1)
    for block in content.replace("\r\n", "\n").split("\n"):
        normalized = block.strip()
        if normalized:
            document.add_paragraph(normalized)
    buffer = BytesIO()
    document.save(buffer)
    return base64.b64encode(buffer.getvalue()).decode("ascii")


def _analysis_prompt(reference_name: str, reference_text: str | None, requested_name: str | None) -> str:
    reference_excerpt = (reference_text or "").strip()[:14000]
    source_section = (
        f"Conteudo extraido do documento:\n{reference_excerpt}"
        if reference_excerpt
        else "A referencia esta em imagem. Analise visualmente campos, titulos, secoes, tabelas e instrucoes do documento."
    )
    requested_name_section = f"Nome sugerido pelo usuario: {requested_name}" if requested_name else "Nome sugerido pelo usuario: nao informado"
    return "\n\n".join(
        [
            "Voce e um especialista em transformar documentos reais em modelos reutilizaveis de preenchimento.",
            "Analise a referencia e crie um modelo de documento para ser preenchido depois por audio transcrito, texto ou por um formulario com campos.",
            requested_name_section,
            f"Arquivo de referencia: {reference_name}",
            source_section,
            "Responda somente JSON valido, sem markdown, neste formato:",
            json.dumps(
                {
                    "name": "Nome curto do modelo",
                    "description": "Para que esse modelo serve",
                    "category": "Categoria",
                    "base_prompt": "Instrucoes para a IA preencher o documento com seguranca",
                    "example_output": "# Titulo do documento\n\n## Campo\n{{valor a preencher}}",
                    "complementary_instructions": "Regras extras, campos obrigatorios e como lidar com dados ausentes",
                    "form_fields": [
                        {
                            "key": "nome_cliente",
                            "label": "Nome do cliente",
                            "type": "text",
                            "placeholder": "Ex.: Joao da Silva",
                            "required": True,
                            "help": "Nome completo que aparece no documento",
                        }
                    ],
                    "output_format": "markdown",
                },
                ensure_ascii=False,
            ),
            "Use placeholders como {{nome_cliente}}, {{data}} ou {{resumo}} quando fizer sentido. Os placeholders devem combinar com as 'key' dos form_fields.",
            "Para cada placeholder que precisa ser preenchido pela pessoa, crie um item em 'form_fields' com key (minusculas, sem acento, underscore no lugar de espaco), label amigavel em portugues, type (text|textarea|date|number), required (true/false) e help opcional.",
            "Nao invente dados do documento final; descreva a estrutura que deve ser preenchida.",
        ]
    )


_FIELD_KEY_PATTERN = re.compile(r"[^a-z0-9_]+")
_PLACEHOLDER_PATTERN = re.compile(r"\{\{\s*([^}]+?)\s*\}\}")


def _slugify_field_key(raw: str) -> str:
    simplified = raw.strip().lower()
    simplified = simplified.replace("ç", "c").replace("ã", "a").replace("á", "a").replace("à", "a")
    simplified = simplified.replace("â", "a").replace("é", "e").replace("ê", "e").replace("í", "i")
    simplified = simplified.replace("ó", "o").replace("ô", "o").replace("õ", "o").replace("ú", "u")
    simplified = simplified.replace("ü", "u").replace(" ", "_").replace("-", "_")
    simplified = _FIELD_KEY_PATTERN.sub("_", simplified).strip("_")
    return simplified[:80] or "campo"


def extract_placeholder_keys(example_output: str | None) -> list[str]:
    if not example_output:
        return []
    seen: list[str] = []
    for match in _PLACEHOLDER_PATTERN.finditer(example_output):
        raw = match.group(1)
        key = _slugify_field_key(raw)
        if key and key not in seen:
            seen.append(key)
    return seen


def build_fallback_form_fields(example_output: str | None) -> list[dict] | None:
    keys = extract_placeholder_keys(example_output)
    if not keys:
        return None
    return [
        {
            "key": key,
            "label": key.replace("_", " ").capitalize(),
            "type": "textarea" if len(key) > 24 else "text",
            "placeholder": None,
            "required": False,
            "help": None,
        }
        for key in keys
    ]


def _normalize_form_fields(raw: object, example_output: str | None) -> list[dict] | None:
    if isinstance(raw, list):
        normalized: list[dict] = []
        seen_keys: set[str] = set()
        for item in raw:
            if not isinstance(item, dict):
                continue
            raw_key = item.get("key") or item.get("label")
            if not isinstance(raw_key, str) or not raw_key.strip():
                continue
            key = _slugify_field_key(raw_key)
            if not key or key in seen_keys:
                continue
            seen_keys.add(key)
            label = item.get("label") if isinstance(item.get("label"), str) and item["label"].strip() else key.replace("_", " ").capitalize()
            field_type = item.get("type") if item.get("type") in {"text", "textarea", "date", "number"} else "text"
            placeholder = item.get("placeholder") if isinstance(item.get("placeholder"), str) and item["placeholder"].strip() else None
            help_text = item.get("help") if isinstance(item.get("help"), str) and item["help"].strip() else None
            required = bool(item.get("required"))
            normalized.append(
                {
                    "key": key,
                    "label": label.strip()[:160],
                    "type": field_type,
                    "placeholder": placeholder,
                    "required": required,
                    "help": help_text,
                }
            )
        if normalized:
            return normalized
    return build_fallback_form_fields(example_output)


def _extract_json_object(raw: str) -> dict[str, object]:
    cleaned = raw.strip()
    if cleaned.startswith("```"):
        cleaned = re.sub(r"^```(?:json)?", "", cleaned, flags=re.IGNORECASE).strip()
        cleaned = re.sub(r"```$", "", cleaned).strip()

    try:
        parsed = json.loads(cleaned)
    except json.JSONDecodeError:
        match = re.search(r"\{.*\}", cleaned, flags=re.DOTALL)
        if not match:
            raise ValueError("A IA nao retornou um JSON de modelo valido") from None
        parsed = json.loads(match.group(0))

    if not isinstance(parsed, dict):
        raise ValueError("A IA retornou um formato inesperado")
    return parsed


def _clean_optional_text(value: object, min_length: int = 1) -> str | None:
    if not isinstance(value, str):
        return None
    normalized = value.strip()
    return normalized if len(normalized) >= min_length else None


def _fallback_payload(
    *,
    filename: str,
    reference_text: str | None,
    name: str | None,
    description: str | None,
    category: str | None,
) -> ReportTemplateCreate:
    inferred_name = name or f"Modelo {Path(filename).stem or datetime.utcnow().strftime('%Y%m%d')}"
    excerpt = (reference_text or "").strip()[:5000]
    example_output = excerpt if len(excerpt) >= 10 else "# Documento\n\n## Campo principal\n{{preencher com base no audio ou texto}}"
    return ReportTemplateCreate(
        name=inferred_name[:120],
        description=(description or "Modelo criado a partir de uma referencia enviada.").strip()[:255],
        category=(category or "Formulario").strip()[:100],
        base_prompt=(
            "Preencha o modelo usando somente as informacoes fornecidas pelo audio transcrito ou texto. "
            "Mantenha a estrutura do documento de referencia e marque como 'Nao informado' quando faltar dado."
        ),
        example_output=example_output,
        complementary_instructions="Preserve titulos, ordem das secoes e campos obrigatorios do documento de referencia.",
        form_fields=build_fallback_form_fields(example_output),
        output_format="markdown",
        is_favorite=False,
    )


def _payload_from_ai_response(
    *,
    raw_response: str,
    filename: str,
    reference_text: str | None,
    name: str | None,
    description: str | None,
    category: str | None,
) -> ReportTemplateCreate:
    data = _extract_json_object(raw_response)
    fallback = _fallback_payload(
        filename=filename,
        reference_text=reference_text,
        name=name,
        description=description,
        category=category,
    )
    example_output = _clean_optional_text(data.get("example_output"), 10) or fallback.example_output
    form_fields = _normalize_form_fields(data.get("form_fields"), example_output)
    return ReportTemplateCreate(
        name=(name or _clean_optional_text(data.get("name"), 3) or fallback.name)[:120],
        description=(description or _clean_optional_text(data.get("description"), 3) or fallback.description)[:255],
        category=(category or _clean_optional_text(data.get("category"), 2) or fallback.category)[:100],
        base_prompt=_clean_optional_text(data.get("base_prompt"), 10) or fallback.base_prompt,
        example_output=example_output,
        complementary_instructions=_clean_optional_text(data.get("complementary_instructions")) or fallback.complementary_instructions,
        form_fields=form_fields,
        output_format="text" if data.get("output_format") == "text" else "markdown",
        is_favorite=False,
    )


def _analyze_reference_payload(
    db: Session,
    reference_file: UploadFile,
    *,
    name: str | None = None,
    description: str | None = None,
    category: str | None = None,
) -> tuple[ReportTemplateCreate, str, str, str | None]:
    filename = reference_file.filename or "referencia"
    content_type = reference_file.content_type or mimetypes.guess_type(filename)[0]
    data = _read_reference_file(reference_file)
    source_format = Path(filename).suffix.lower().lstrip(".") or (content_type or "arquivo")
    is_image = _is_image_reference(filename, content_type)
    reference_text = None if is_image else _extract_reference_text(filename, content_type, data)
    prompt = _analysis_prompt(filename, reference_text, name)

    if is_image:
        raw_analysis = _generate_reference_analysis_image(db, prompt, data, content_type)
    else:
        raw_analysis = _generate_reference_analysis_text(db, prompt)

    payload = (
        _payload_from_ai_response(
            raw_response=raw_analysis,
            filename=filename,
            reference_text=reference_text,
            name=name,
            description=description,
            category=category,
        )
        if raw_analysis
        else _fallback_payload(
            filename=filename,
            reference_text=reference_text,
            name=name,
            description=description,
            category=category,
        )
    )
    return payload, filename, source_format, reference_text


def analyze_template_reference(
    db: Session,
    reference_file: UploadFile,
    *,
    name: str | None = None,
    description: str | None = None,
    category: str | None = None,
) -> ReportTemplateReferenceAnalysis:
    payload, filename, source_format, reference_text = _analyze_reference_payload(
        db,
        reference_file,
        name=name,
        description=description,
        category=category,
    )
    converted_docx_filename = None
    converted_docx_base64 = None
    if source_format == "pdf" and reference_text:
        converted_docx_filename = f"{Path(filename).stem or 'documento'}-convertido.docx"
        converted_docx_base64 = _build_docx_base64_from_text(Path(filename).stem or "Documento convertido", reference_text)

    return ReportTemplateReferenceAnalysis(
        **payload.model_dump(),
        source_filename=filename,
        source_format=source_format,
        converted_docx_filename=converted_docx_filename,
        converted_docx_base64=converted_docx_base64,
    )


def extract_template_reference_text(reference_file: UploadFile) -> ReportTemplateReferenceText:
    filename = reference_file.filename or "referencia"
    content_type = reference_file.content_type or mimetypes.guess_type(filename)[0]
    if _is_image_reference(filename, content_type):
        raise ValueError("Imagens precisam da analise por IA; nao foi possivel abrir como texto.")

    data = _read_reference_file(reference_file)
    source_format = Path(filename).suffix.lower().lstrip(".") or (content_type or "arquivo")
    content = _extract_reference_text(filename, content_type, data)
    return ReportTemplateReferenceText(
        source_filename=filename,
        source_format=source_format,
        content=content,
    )


def _generate_reference_analysis_text(db: Session, prompt: str) -> str | None:
    settings_data = get_effective_provider_settings(db)
    for provider in _report_provider_order(settings_data):
        try:
            if provider == "openai" and isinstance(settings_data.get("openai_api_key"), str) and settings_data["openai_api_key"]:
                content, _ = _generate_openai(prompt, str(settings_data["openai_api_key"]))
                return content
            if provider == "claude" and isinstance(settings_data.get("claude_api_key"), str) and settings_data["claude_api_key"]:
                content, _ = _generate_claude(prompt, str(settings_data["claude_api_key"]))
                return content
            if provider == "gemini" and isinstance(settings_data.get("gemini_api_key"), str) and settings_data["gemini_api_key"]:
                content, _ = _generate_gemini(prompt, str(settings_data["gemini_api_key"]))
                return content
        except Exception:
            continue
    return None


def _generate_reference_analysis_image(db: Session, prompt: str, data: bytes, content_type: str | None) -> str:
    settings_data = get_effective_provider_settings(db)
    openai_key = settings_data.get("openai_api_key")
    if not isinstance(openai_key, str) or not openai_key:
        raise ValueError("Configure uma chave OpenAI para a IA analisar imagens de modelos.")

    from openai import OpenAI

    media_type = content_type or "image/png"
    client = OpenAI(api_key=openai_key, timeout=get_settings().provider_timeout_seconds)
    encoded = base64.b64encode(data).decode("ascii")
    response = client.chat.completions.create(
        model=get_settings().openai_report_model,
        messages=[
            {"role": "system", "content": "Voce transforma documentos e imagens em modelos reutilizaveis de formulario."},
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": prompt},
                    {"type": "image_url", "image_url": {"url": f"data:{media_type};base64,{encoded}"}},
                ],
            },
        ],
    )
    content = response.choices[0].message.content or ""
    if not content.strip():
        raise ValueError("A IA retornou uma analise vazia da imagem.")
    return content.strip()


def _unique_template_name(repository: ReportTemplateRepository, requested_name: str, workspace_id: str) -> str:
    normalized = requested_name.strip()[:120] or "Modelo gerado por IA"
    if not repository.get_by_name(normalized, workspace_id):
        return normalized

    suffix = 2
    while True:
        candidate = f"{normalized[:110]} {suffix}"
        if not repository.get_by_name(candidate, workspace_id):
            return candidate
        suffix += 1


def create_template_from_reference(
    db: Session,
    reference_file: UploadFile,
    *,
    name: str | None = None,
    description: str | None = None,
    category: str | None = None,
    workspace_id: str = "local-workspace",
) -> ReportTemplate:
    payload, _, _, _ = _analyze_reference_payload(
        db,
        reference_file,
        name=name,
        description=description,
        category=category,
    )

    repository = ReportTemplateRepository(db)
    payload.name = _unique_template_name(repository, payload.name, workspace_id)
    return repository.create(ReportTemplate(workspace_id=workspace_id, **payload.model_dump()))


def update_template(db: Session, template_id: str, payload: ReportTemplateUpdate, workspace_id: str = "local-workspace") -> ReportTemplate:
    repository = ReportTemplateRepository(db)
    template = repository.get_for_workspace(template_id, workspace_id)
    if not template:
        raise ValueError("Modelo não encontrado")

    new_name = payload.name.strip() if isinstance(payload.name, str) else None
    if new_name and new_name != template.name:
        existing = repository.get_by_name(new_name, workspace_id)
        if existing and existing.id != template.id:
            raise ValueError("Já existe um modelo com esse nome")

    for field, value in payload.model_dump(exclude_unset=True).items():
        setattr(template, field, value)
    return repository.save(template)


def delete_template(db: Session, template_id: str, workspace_id: str = "local-workspace") -> None:
    repository = ReportTemplateRepository(db)
    template = repository.get_for_workspace(template_id, workspace_id)
    if not template:
        raise ValueError("Modelo não encontrado")
    repository.delete(template)


def duplicate_template(db: Session, template_id: str, workspace_id: str = "local-workspace") -> ReportTemplate:
    repository = ReportTemplateRepository(db)
    template = repository.get_for_workspace(template_id, workspace_id)
    if not template:
        raise ValueError("Modelo não encontrado")
    duplicate = ReportTemplate(
        workspace_id=workspace_id,
        name=_build_duplicate_name(repository, template.name, workspace_id),
        description=template.description,
        category=template.category,
        base_prompt=template.base_prompt,
        example_output=template.example_output,
        complementary_instructions=template.complementary_instructions,
        form_fields=template.form_fields,
        output_format=template.output_format,
        is_favorite=False,
    )
    return repository.create(duplicate)
