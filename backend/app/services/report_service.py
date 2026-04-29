import re
from dataclasses import dataclass
from pathlib import Path
from xml.sax.saxutils import escape

from sqlalchemy.orm import Session

from app.core.config import get_settings
from app.models.enums import ProcessingStatus, ReportFormat, TranscriptionEngine
from app.models.report import GeneratedReport
from app.repositories.report_repository import ReportRepository
from app.repositories.report_template_repository import ReportTemplateRepository
from app.repositories.upload_repository import UploadRepository
from app.schemas.report import GenerateReportRequest, ReportExportExtension
from app.services.settings_service import get_effective_provider_settings


@dataclass(frozen=True)
class ReportExportFile:
    extension: ReportExportExtension
    path: Path
    media_type: str
    filename: str


REPORT_EXPORT_MEDIA_TYPES = {
    ReportExportExtension.MD: "text/markdown; charset=utf-8",
    ReportExportExtension.TXT: "text/plain; charset=utf-8",
    ReportExportExtension.DOCX: "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ReportExportExtension.PDF: "application/pdf",
}
EXPORT_CONTROL_CHARACTERS = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]")


def _markdown_heading_level(line: str) -> int:
    stripped = line.lstrip()
    if not stripped.startswith("#"):
        return 0
    level = len(stripped) - len(stripped.lstrip("#"))
    if level <= 0 or (len(stripped) > level and stripped[level] != " "):
        return 0
    return min(level, 6)


def _sanitize_export_text(value: str) -> str:
    normalized = value.replace("\r\n", "\n").replace("\r", "\n")
    return EXPORT_CONTROL_CHARACTERS.sub("", normalized)


def write_docx_export(export_path: Path, content: str) -> None:
    from docx import Document

    sanitized_content = _sanitize_export_text(content)
    document = Document()
    for line in sanitized_content.split("\n"):
        normalized_line = line.rstrip()
        if not normalized_line:
            document.add_paragraph("")
            continue

        heading_level = _markdown_heading_level(normalized_line)
        if heading_level:
            document.add_heading(normalized_line[heading_level:].strip(), level=min(heading_level, 4))
            continue

        document.add_paragraph(normalized_line)

    document.save(export_path)


def write_pdf_export(export_path: Path, title: str, content: str) -> None:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    sanitized_title = _sanitize_export_text(title)
    sanitized_content = _sanitize_export_text(content)
    styles = getSampleStyleSheet()
    body_style = ParagraphStyle(
        "ReportBody",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=10,
        leading=14,
        spaceAfter=10,
    )
    heading_styles = {
        1: ParagraphStyle("Heading1Export", parent=styles["Heading1"], fontName="Helvetica-Bold"),
        2: ParagraphStyle("Heading2Export", parent=styles["Heading2"], fontName="Helvetica-Bold"),
        3: ParagraphStyle("Heading3Export", parent=styles["Heading3"], fontName="Helvetica-Bold"),
        4: ParagraphStyle("Heading4Export", parent=styles["Heading4"], fontName="Helvetica-Bold"),
    }

    document = SimpleDocTemplate(str(export_path), pagesize=A4, title=sanitized_title)
    story = []
    for block in sanitized_content.split("\n\n"):
        normalized_block = block.strip()
        if not normalized_block:
            continue

        heading_level = _markdown_heading_level(normalized_block)
        if heading_level:
            story.append(Paragraph(escape(normalized_block[heading_level:].strip()), heading_styles[min(heading_level, 4)]))
            story.append(Spacer(1, 8))
            continue

        paragraph_html = "<br/>".join(escape(line) for line in normalized_block.split("\n"))
        story.append(Paragraph(paragraph_html, body_style))
        story.append(Spacer(1, 6))

    document.build(story)


def _write_report_exports(export_dir: Path, report_id: str, title: str, output_format: ReportFormat, content: str) -> None:
    primary_suffix = ".md" if output_format == ReportFormat.MARKDOWN else ".txt"
    primary_path = export_dir / f"{report_id}{primary_suffix}"
    primary_path.write_text(content, encoding="utf-8")
    write_docx_export(export_dir / f"{report_id}.docx", content)
    write_pdf_export(export_dir / f"{report_id}.pdf", title, content)


def build_report_prompt(
    transcription: str,
    template_prompt: str | None,
    example_output: str | None,
    custom_request: str | None,
    additional_instructions: str | None,
) -> str:
    sections = [
        "Você vai gerar um relatório estruturado a partir de uma transcrição.",
        "Use apenas informações presentes na transcrição. Não invente fatos, nomes, datas ou números.",
        "Se o modelo pedir um campo ausente, escreva 'Não informado na transcrição'.",
        f"Transcrição base:\n{transcription}",
    ]
    if template_prompt:
        sections.append(f"Objetivo do modelo:\n{template_prompt}")
    if example_output:
        sections.append(
            "Modelo de referência para ser analisado e seguido na estrutura final:\n"
            f"{example_output}\n\n"
            "Copie a organização, a ordem das seções e o estilo dos títulos, mas preencha o conteúdo usando somente a transcrição."
        )
    if custom_request:
        sections.append(f"Pedido do usuário:\n{custom_request}")
    if additional_instructions:
        sections.append(f"Instruções adicionais:\n{additional_instructions}")
    sections.append("Entregue o relatório final pronto para uso, já preenchido com base na transcrição.")
    return "\n\n".join(sections)


def _generate_openai(prompt: str, api_key: str) -> tuple[str, TranscriptionEngine]:
    from openai import OpenAI

    settings = get_settings()
    client = OpenAI(api_key=api_key, timeout=settings.provider_timeout_seconds)
    response = client.chat.completions.create(
        model=settings.openai_report_model,
        messages=[
            {"role": "system", "content": "Você gera relatórios claros, objetivos e bem estruturados."},
            {"role": "user", "content": prompt},
        ],
    )
    content = response.choices[0].message.content or ""
    if not content.strip():
        raise RuntimeError("OpenAI retornou relatório vazio")
    return content.strip(), TranscriptionEngine.OPENAI


def _generate_gemini(prompt: str, api_key: str) -> tuple[str, TranscriptionEngine]:
    from google import genai

    settings = get_settings()
    client = genai.Client(api_key=api_key)
    response = client.models.generate_content(model=settings.gemini_report_model, contents=prompt)
    content = getattr(response, "text", None) or ""
    if not content.strip():
        raise RuntimeError("Gemini retornou relatório vazio")
    return content.strip(), TranscriptionEngine.GEMINI


def _generate_claude(prompt: str, api_key: str) -> tuple[str, TranscriptionEngine]:
    from anthropic import Anthropic

    settings = get_settings()
    client = Anthropic(api_key=api_key, timeout=settings.provider_timeout_seconds)
    response = client.messages.create(
        model=settings.claude_report_model,
        max_tokens=4096,
        system="Você gera relatórios claros, objetivos e bem estruturados.",
        messages=[{"role": "user", "content": prompt}],
    )
    blocks = [getattr(block, "text", "") for block in getattr(response, "content", [])]
    content = "\n".join(block for block in blocks if block).strip()
    if not content:
        raise RuntimeError("Claude retornou relatório vazio")
    return content, TranscriptionEngine.CLAUDE


def _generate_local_fallback(title: str, transcription: str, custom_request: str | None) -> tuple[str, TranscriptionEngine]:
    excerpt = transcription[:3000].strip()
    content = (
        f"# {title}\n\n"
        "## Solicitação\n"
        f"{custom_request or 'Relatório baseado na transcrição.'}\n\n"
        "## Resumo inicial\n"
        f"{excerpt}\n"
    )
    return content, TranscriptionEngine.NONE


def _report_provider_order(settings_data: dict[str, str | int | bool | None]) -> list[str]:
    raw_order = settings_data.get("report_provider_order")
    if isinstance(raw_order, list):
        normalized = [str(item).strip().lower() for item in raw_order]
        ordered = [item for item in normalized if item in {"openai", "claude", "gemini", "local"}]
        return ordered or ["openai", "claude", "gemini", "local"]

    return ["openai", "claude", "gemini", "local"]


def _read_report_or_error(db: Session, report_id: str, workspace_id: str | None = None) -> GeneratedReport:
    repository = ReportRepository(db)
    report = repository.get_for_workspace(report_id, workspace_id) if workspace_id else repository.get(report_id)
    if not report:
        raise ValueError("RelatÃ³rio nÃ£o encontrado")
    return report


def _report_exports_dir(db: Session) -> Path:
    settings_data = get_effective_provider_settings(db)
    return Path(str(settings_data.get("export_directory") or get_settings().exports_dir))


def build_export_download_filename(title: str, extension: ReportExportExtension, fallback: str = "relatorio") -> str:
    normalized = re.sub(r"[\\/:*?\"<>|]+", "-", title).strip()
    normalized = re.sub(r"\s+", " ", normalized)[:120]
    return f"{normalized or fallback}.{extension.value}"


def _report_download_filename(title: str, extension: ReportExportExtension) -> str:
    return build_export_download_filename(title, extension)


def _expected_report_export_extensions(report: GeneratedReport) -> list[ReportExportExtension]:
    primary_extension = ReportExportExtension.MD if report.output_format == ReportFormat.MARKDOWN else ReportExportExtension.TXT
    return [primary_extension, ReportExportExtension.DOCX, ReportExportExtension.PDF]


def _report_export_candidates(report: GeneratedReport, export_dir: Path) -> list[ReportExportFile]:
    return [
        ReportExportFile(
            extension=extension,
            path=export_dir / f"{report.id}.{extension.value}",
            media_type=REPORT_EXPORT_MEDIA_TYPES[extension],
            filename=_report_download_filename(report.title, extension),
        )
        for extension in _expected_report_export_extensions(report)
    ]


def list_report_exports(db: Session, report_id: str, workspace_id: str | None = None) -> list[ReportExportFile]:
    report = _read_report_or_error(db, report_id, workspace_id)
    export_dir = _report_exports_dir(db)
    return [export_file for export_file in _report_export_candidates(report, export_dir) if export_file.path.is_file()]


def get_report_export(db: Session, report_id: str, extension: ReportExportExtension, workspace_id: str | None = None) -> ReportExportFile:
    report = _read_report_or_error(db, report_id, workspace_id)
    export_dir = _report_exports_dir(db)

    for export_file in _report_export_candidates(report, export_dir):
        if export_file.extension != extension:
            continue
        if not export_file.path.is_file():
            raise ValueError("ExportaÃ§Ã£o nÃ£o encontrada")
        return export_file

    raise ValueError("Formato de exportaÃ§Ã£o nÃ£o disponÃ­vel para este relatÃ³rio")


def rename_report(db: Session, report_id: str, title: str, workspace_id: str | None = None) -> GeneratedReport:
    report_repository = ReportRepository(db)
    report = report_repository.get_for_workspace(report_id, workspace_id) if workspace_id else report_repository.get(report_id)
    if not report:
        raise ValueError("Relatório não encontrado")

    report.title = title.strip()
    return report_repository.save(report)


def generate_report(db: Session, payload: GenerateReportRequest, workspace_id: str = "local-workspace") -> GeneratedReport:
    upload_repository = UploadRepository(db)
    template_repository = ReportTemplateRepository(db)
    report_repository = ReportRepository(db)
    upload = upload_repository.get_for_workspace(payload.upload_id, workspace_id)
    if not upload:
        raise ValueError("Upload não encontrado")
    if not upload.transcription_text:
        raise ValueError("A transcrição ainda não está disponível")

    template = template_repository.get_for_workspace(payload.template_id, workspace_id) if payload.template_id else None
    output_format = template.output_format if template else ReportFormat.MARKDOWN
    prompt = build_report_prompt(
        transcription=upload.transcription_text,
        template_prompt=template.base_prompt if template else None,
        example_output=template.example_output if template else None,
        custom_request=payload.custom_request,
        additional_instructions=payload.additional_instructions or (template.complementary_instructions if template else None),
    )

    upload.status = ProcessingStatus.GENERATING_REPORT
    upload_repository.save(upload)

    settings_data = get_effective_provider_settings(db)

    content: str | None = None
    engine = TranscriptionEngine.NONE
    for provider in _report_provider_order(settings_data):
        try:
            if provider == "openai":
                openai_key = settings_data.get("openai_api_key")
                if isinstance(openai_key, str) and openai_key:
                    content, engine = _generate_openai(prompt, openai_key)
                    break
                continue

            if provider == "claude":
                claude_key = settings_data.get("claude_api_key")
                if isinstance(claude_key, str) and claude_key:
                    content, engine = _generate_claude(prompt, claude_key)
                    break
                continue

            if provider == "gemini":
                gemini_key = settings_data.get("gemini_api_key")
                if isinstance(gemini_key, str) and gemini_key:
                    content, engine = _generate_gemini(prompt, gemini_key)
                    break
                continue

            if provider == "local":
                content, engine = _generate_local_fallback(payload.title, upload.transcription_text, payload.custom_request)
                break
        except Exception:
            continue

    if content is None:
        content, engine = _generate_local_fallback(payload.title, upload.transcription_text, payload.custom_request)

    report = GeneratedReport(
        workspace_id=workspace_id,
        upload_id=upload.id,
        template_id=template.id if template else None,
        title=payload.title,
        request_prompt=prompt,
        content=content,
        output_format=output_format,
        generator_engine=engine,
    )
    created = report_repository.create(report)
    upload.report_count += 1
    upload.status = ProcessingStatus.COMPLETED
    upload_repository.save(upload)

    export_dir = Path(str(settings_data.get("export_directory") or get_settings().exports_dir))
    export_dir.mkdir(parents=True, exist_ok=True)
    _write_report_exports(export_dir, created.id, created.title, output_format, created.content)

    return created
