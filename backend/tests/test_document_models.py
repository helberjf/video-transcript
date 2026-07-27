from io import BytesIO

import pytest
from docx import Document
from starlette.datastructures import UploadFile

from app.api.routes.document_models import create_document_model_endpoint
from app.schemas.document_model import DocumentModelCreate
from app.services.document_model_service import create_document_model, list_document_models
from tests.conftest import create_test_session


def _make_docx_upload(filename: str = "modelo.docx") -> UploadFile:
    document = Document()
    document.add_paragraph("Cliente: Maria Silva")
    document.add_paragraph("Contexto: reunião comercial")
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return UploadFile(file=buffer, filename=filename)


def test_create_document_model_from_docx(tmp_path, monkeypatch) -> None:
    session, engine = create_test_session(tmp_path)
    storage_dir = tmp_path / "storage"
    monkeypatch.setattr(
        "app.services.document_model_service.get_settings",
        lambda: type("Settings", (), {"storage_dir": storage_dir})(),
    )

    created = create_document_model(
        session,
        _make_docx_upload(),
        name="Modelo de ata",
        description="Ata base para reuniões",
        category="reuniao",
        default_context="Tom formal e objetivo",
    )

    listed = list_document_models(session)

    assert created.name == "Modelo de ata"
    assert created.description == "Ata base para reuniões"
    assert created.category == "reuniao"
    assert created.default_context == "Tom formal e objetivo"
    assert created.source_filename == "modelo.docx"
    assert "Cliente: Maria Silva" in created.source_text
    assert created.source_path.endswith("modelo.docx")
    assert len(listed) == 1

    session.close()
    engine.dispose()


def test_create_document_model_rejects_duplicate_name(tmp_path, monkeypatch) -> None:
    session, engine = create_test_session(tmp_path)
    storage_dir = tmp_path / "storage"
    monkeypatch.setattr(
        "app.services.document_model_service.get_settings",
        lambda: type("Settings", (), {"storage_dir": storage_dir})(),
    )

    create_document_model(
        session,
        _make_docx_upload("modelo-1.docx"),
        name="Modelo padrão",
        description="Primeiro modelo",
        category="documento",
        default_context="Usar linguagem clara",
    )

    with pytest.raises(ValueError, match="Já existe um modelo de documento com esse nome"):
        create_document_model(
            session,
            _make_docx_upload("modelo-2.docx"),
            name="Modelo padrão",
            description="Segundo modelo",
            category="documento",
            default_context="Usar linguagem clara",
        )

    session.close()
    engine.dispose()


def test_document_model_api_creates_record(tmp_path, monkeypatch) -> None:
    session, engine = create_test_session(tmp_path)
    storage_dir = tmp_path / "storage"
    monkeypatch.setattr(
        "app.services.document_model_service.get_settings",
        lambda: type("Settings", (), {"storage_dir": storage_dir})(),
    )

    response = create_document_model_endpoint(
        file=UploadFile(file=BytesIO("Cliente: Maria Silva\nContexto: reunião comercial".encode("utf-8")), filename="modelo.txt"),
        name="Modelo via API",
        description="Criado pela API",
        category="documento",
        default_context="Contexto padrão salvo",
        base_instructions=None,
        db=session,
        workspace_id="local-workspace",
    )

    assert response.name == "Modelo via API"
    assert response.default_context == "Contexto padrão salvo"
    assert response.source_filename == "modelo.txt"

    session.close()
    engine.dispose()